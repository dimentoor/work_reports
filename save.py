import pandas as pd
import json
import database
import docx
from io import BytesIO
from docx.shared import Inches


class ExcelLoader:
    def __init__(self, open_path, sheet_name):
        self.open_path = open_path
        self.sheet_name = sheet_name
        self.table = pd.DataFrame()

    # open file
    def open_file(self):
        self.table = pd.read_excel(self.open_path, sheet_name=self.sheet_name)
        print("Opened {}.".format(self.open_path))


class ExcelDumper:
    def __init__(self, save_path):
        self.save_path = save_path

    # write file
    @staticmethod
    def write_file(filename, samples: dict):
        with pd.ExcelWriter(filename) as writer:
            for sample_name, sample in samples.items():
                sample.to_excel(writer, sheet_name=sample_name)
        print("Wrote to {}.".format(filename))


class WordDumper:
    def __init__(self, save_path):
        self.save_path = save_path

    # write file
    @staticmethod
    def write_file(filename, samples: dict):
        document = docx.Document()

        for key, value in samples.items():
            document.add_heading(key, level=1)

            if isinstance(value, pd.DataFrame):
                if not value.empty:
                    # Создаем таблицу
                    table = document.add_table(rows=value.shape[0] + 1, cols=value.shape[1])
                    # Устанавливаем стиль 'Table Grid' для отображения границ
                    table.style = 'Table Grid'
                    # Добавляем заголовки столбцов
                    for col_num, col_name in enumerate(value.columns):
                        table.cell(0, col_num).text = col_name
                    # Добавляем данные в таблицу
                    for row_num in range(value.shape[0]):
                        for col_num in range(value.shape[1]):
                            table.cell(row_num + 1, col_num).text = str(value.iloc[row_num, col_num])
                else:
                    document.add_paragraph("Данные приведены в excel файле")
            elif isinstance(value, BytesIO):
                document.add_picture(value, width=Inches(4))
            else:
                document.add_paragraph(str(value))

        document.save(filename)
        print("Wrote to {}.".format(filename))


class MongoLoader:
    pass


class MongoDumper:
    def __init__(self):
        pass

    @staticmethod
    def df_to_json(class_object, collection_name):
        collection = database.Mongo.create_collection(collection_name)

        for sample_name, sample in class_object.items():
            # print(type(sample))
            tmp = sample.to_json(force_ascii=False, orient="index", indent=4)
            # print(type(tmp))
            # print(json.loads(tmp))
            # print('\n')
            collection_id = database.Mongo.add_doc(collection, json.loads(tmp))

        print("ADD to database {}.".format(collection_name))


# with open('{}.json'.format(sample_name), 'w') as outfile:
#     json.dump(tmp, outfile, ensure_ascii=False, indent=4, sort_keys=True)
